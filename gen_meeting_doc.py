from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

RED   = RGBColor(0xC0, 0x39, 0x2B)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF9, 0xF9, 0xF9)

LOGO_PATH = r"C:\Users\User\Downloads\Myke_Agent\Dentstage凹痕工廠相關事項\Dent Stage LOGO 圓 更新.png"
OUT_PATH  = r"C:\Users\User\Downloads\Myke_Agent\Dentstage凹痕工廠相關事項\會議紀錄_20260622.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(side, {'val':'nil'})
        el  = OxmlElement(f'w:{side}')
        for k,v in val.items():
            el.set(qn(f'w:{k}'), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def run(para, text, bold=False, color=None, size=11, font='Microsoft JhengHei'):
    r = para.add_run(text)
    r.bold = bold
    r.font.name = font
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        run(p, text, bold=True, color=BLACK, size=13)
        p.paragraph_format.border_bottom = True
        # red underline via bottom border on paragraph
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        btm  = OxmlElement('w:bottom')
        btm.set(qn('w:val'),  'single')
        btm.set(qn('w:sz'),   '6')
        btm.set(qn('w:space'),'1')
        btm.set(qn('w:color'),'C0392B')
        pBdr.append(btm)
        pPr.append(pBdr)
    else:
        run(p, text, bold=True, color=RED, size=11.5)
    return p

def body_para(text, bullet=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if bullet:
        p.paragraph_format.left_indent   = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        run(p, '▸  ', bold=False, color=RED, size=10.5)
        run(p, text, size=10.5)
    elif indent:
        p.paragraph_format.left_indent = Cm(1.2)
        run(p, text, size=10.5, color=RGBColor(0x44,0x44,0x44))
    else:
        run(p, text, size=10.5)
    return p

def make_table(headers, rows, col_widths=None, header_bg='1A1A1A'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # header row
    hr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run(p, h, bold=True, color=WHITE, size=10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, 'F9F9F9')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run(p, val, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# HEADER TABLE (Logo + Title + Date)
# ══════════════════════════════════════════════════════════════════════════════
ht = doc.add_table(rows=1, cols=3)
ht.alignment = WD_TABLE_ALIGNMENT.CENTER

logo_cell  = ht.rows[0].cells[0]
title_cell = ht.rows[0].cells[1]
date_cell  = ht.rows[0].cells[2]

# background
for c in (logo_cell, title_cell, date_cell):
    set_cell_bg(c, '1A1A1A')
    set_cell_border(c,
        top={'val':'nil'}, left={'val':'nil'},
        bottom={'val':'nil'}, right={'val':'nil'})

# logo
logo_cell.width = Cm(3.2)
logo_p = logo_cell.paragraphs[0]
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_before = Pt(6)
logo_p.paragraph_format.space_after  = Pt(6)
logo_r = logo_p.add_run()
logo_r.add_picture(LOGO_PATH, width=Cm(2.4))

# title
title_cell.width = Cm(10)
title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
t1 = title_cell.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.LEFT
t1.paragraph_format.space_before = Pt(4)
run(t1, 'DENTSTAGE 凹痕工廠', bold=True, color=WHITE, size=15)
t2 = title_cell.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.LEFT
t2.paragraph_format.space_after = Pt(4)
run(t2, '部門工作分配與業務流程規劃會議', color=RGBColor(0xC0,0xC0,0xC0), size=10)

# date
date_cell.width = Cm(3.5)
date_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
d1 = date_cell.add_paragraph()
d1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
d1.paragraph_format.space_before = Pt(4)
run(d1, '2026.06.22', bold=True, color=RED, size=16)
d2 = date_cell.add_paragraph()
d2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
d2.paragraph_format.space_after = Pt(4)
run(d2, 'MEETING MINUTES', color=RGBColor(0x88,0x88,0x88), size=8)

doc.add_paragraph()

# ── Meta info bar ─────────────────────────────────────────────────────────────
mt = doc.add_table(rows=1, cols=4)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [('主持', 'Jay'), ('出席', 'Jay、Mary'), ('地點', 'Dentstage 凹痕工廠'), ('性質', '內部工作分配會議')]
for i, (label, val) in enumerate(meta):
    c = mt.rows[0].cells[i]
    set_cell_bg(c, 'C0392B')
    set_cell_border(c, top={'val':'nil'}, left={'val':'nil'}, bottom={'val':'nil'}, right={'val':'nil'})
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f'{label}：', bold=True, color=RGBColor(0xFF,0xDD,0xDD), size=9)
    run(p, val, color=WHITE, size=9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 人員職責分工
# ══════════════════════════════════════════════════════════════════════════════
heading('1．人員職責分工')
make_table(
    ['工作項目', '負責人', '備註'],
    [
        ['採購（下單、帳號管理）',   'Jay',              'Jay 提供各供應商網站帳號密碼給採購同仁'],
        ['進貨驗收 / 倉管',           '主任',             '7/1 入職後接手；Jay 採購完先 key 採購單，主任對單驗收'],
        ['出貨 / 銷貨',               '主任',             '不在時由Mary cover'],
        ['庫存扣帳（實體倉）',        '主任',             '出貨後執行；網站庫存由系統自動扣'],
        ['報價單出單',                'Jay ＋ 主任',      'Dingxin 上負責人欄位須填本人姓名，不得出現對方名字'],
        ['網站上架 / 下架',           'Jay',              '電腦操作由 Jay 負責'],
        ['網站庫存扣帳',              '主任',             '非網站訂單的手動扣帳由主任執行'],
        ['FB / IG 貼文',              'Jay',              ''],
        ['客訴 / 換貨 / 瑕疵品',     'Jay（主力）',      '國外窗口由 Jay 負責；大陸客戶由主任配合'],
        ['外訪廠商 / 學校',           '主任',             'Jay 外出時由主任代理'],
    ],
    col_widths=[5.5, 3.0, 8.2]
)

body_para('※ 主任到職日：2026 年 7 月 1 日，入職第一週先執行盤點，熟悉倉庫與系統後正式接手。', indent=True)

# ══════════════════════════════════════════════════════════════════════════════
# 2. 採購供應商整理
# ══════════════════════════════════════════════════════════════════════════════
heading('2．採購供應商窗口整理')
make_table(
    ['供應商', '下單方式', '付款 / 備註'],
    [
        ['KECO',                  '網站下單（Jay 提供帳號密碼）',              '天篷（Kimi）付款'],
        ['VIP',                   'FB 群組 / 微信群（Jay 已建群）',             '很久未購入，重新聯繫中'],
        ['Altra',                 '供應商直接下單 → 寄台灣',                   '多品項需下架（已被取代）；大量品項待下價'],
        ['Anson',                 '窗口 Lisa（待確認聯繫方式）',               ''],
        ['Blaine',                '網站下單（B-L-E-H-N）',                     '可寄至 Altra 倉轉運'],
        ['PDR Finesse（PTRFNS）', 'KECO / Dencroft 網站順便購買',             '無折扣；Dencroft 僅少量折扣'],
        ['Dencroft',              '網站下單',                                   '頭（tips）需從網站購買；桿子另外談'],
        ['Brickwick',             'Email 窗口（待補充）',                       '8～9 折優惠，Jay 尚未聯繫，待銜接'],
        ['Stanley / Stanliner',   '暫不開放',                                  '每項商品較特殊，待 Jay 直接處理'],
        ['自製品（金剛等）',       '由 Jay 聯絡廠商',                           'Kenji 光板退款事宜（每片 20–30 RMB）待確認'],
        ['Amazon（錘子）',         '透過 Sally 詢價',                           '見第 6 節討論事項'],
    ],
    col_widths=[3.5, 5.5, 7.7]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. 銷貨流程規範
# ══════════════════════════════════════════════════════════════════════════════
heading('3．銷貨流程規範')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(8)
run(p, '核心原則：收到款項即銷貨（不等出貨）', bold=True, color=RED, size=11)

heading('■ 一般散客（網路下單 / 刷卡 / 匯款）', level=2)
ft = doc.add_table(rows=1, cols=7)
flow_steps = ['客戶下單付款', '→', '確認收款', '→', '銷貨', '→', '出貨']
for i, step in enumerate(flow_steps):
    c = ft.rows[0].cells[i]
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if step == '→':
        set_cell_bg(c, 'FFFFFF')
        run(p, step, bold=True, color=RED, size=14)
    else:
        set_cell_bg(c, '1A1A1A')
        run(p, step, bold=True, color=WHITE, size=9.5)
    set_cell_border(c, top={'val':'nil'}, left={'val':'nil'}, bottom={'val':'nil'}, right={'val':'nil'})
doc.add_paragraph()

heading('■ 特殊情況：先出貨、後收款（學校 / 公司行號）', level=2)
ft2 = doc.add_table(rows=1, cols=9)
flow_steps2 = ['確認訂單', '→', '出貨', '→', '銷貨', '→', '進入代結帳區', '→', '確認收款後結清']
for i, step in enumerate(flow_steps2):
    c = ft2.rows[0].cells[i]
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if step == '→':
        set_cell_bg(c, 'FFFFFF')
        run(p, step, bold=True, color=RED, size=14)
    else:
        set_cell_bg(c, '1A1A1A')
        run(p, step, bold=True, color=WHITE, size=9.5)
    set_cell_border(c, top={'val':'nil'}, left={'val':'nil'}, bottom={'val':'nil'}, right={'val':'nil'})
doc.add_paragraph()
body_para('「代結帳區」由主任管控，包含：客戶名稱、金額、已結/未結狀態，與 Jay / 天篷（Kimi）共用雲端表格。', indent=True)

heading('■ 帳款結算方式', level=2)
for item in [
    '一般帳款（台幣）：直接與天篷（Kimi）結算',
    '主任月帳：與天篷（Kimi）結算',
    '人民幣帳款：與 Jay 結算（使用 Jay 帳戶）',
    '現場預備金：NT$30,000（由天篷（Kimi）每月補回）',
]:
    body_para(item, bullet=True)

# ══════════════════════════════════════════════════════════════════════════════
# 4. 大陸業務對接
# ══════════════════════════════════════════════════════════════════════════════
heading('4．大陸業務對接')
make_table(
    ['項目', '負責人', '說明'],
    [
        ['嫂子（大陸代理）日常對接',   '主任',   '原由 Jay 處理，7/1 後移交主任；嫂子已知情'],
        ['微信客服窗口',               '主任',   '使用公司手機；微信名稱建議改為「DentStage 張[主任姓名]」'],
        ['大陸客戶網站下單',           'Jay',    'Jay 確認後通知主任出貨'],
        ['大陸倉庫存管理',             '主任',   '移倉已完成；主任負責盤點與出貨給嫂子'],
        ['大陸倉售價制定',             '主任',   '主任負責管控成本與利潤，可調整售價'],
        ['WhatsApp / 其他國際客戶',    'Jay',    '非大陸地區國際客戶統一由 Jay 處理'],
    ],
    col_widths=[4.5, 3.0, 9.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. 盤點計畫
# ══════════════════════════════════════════════════════════════════════════════
heading('5．盤點計畫')
for item in [
    '主任入職後第一週執行全倉盤點，作為熟悉倉庫與系統的實作訓練',
    '盤點範圍：Dingxin 大陸倉倉別清單（Jay 已移入完畢，數量最完整）',
    '盤點前需確認中間進銷貨異動，避免數量差異',
    '盤點完成後才可正式提供嫂子（大陸倉）庫存數量',
    '設定半年一次定期盤點制度',
]:
    body_para(item, bullet=True)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run(p, '⚠ 目前倉庫已部分盤點，移倉後有品項未即時銷貨扣帳，建議主任入職後優先清帳，避免庫存落差持續擴大。', bold=True, color=RED, size=10)

# ══════════════════════════════════════════════════════════════════════════════
# 6. 其他討論事項
# ══════════════════════════════════════════════════════════════════════════════
heading('6．其他討論事項')

heading('■ 錘子採購評估（Amazon / Sally）', level=2)
for item in [
    'Sally 詢價結果：100 支 × US$38.5 / 支，折扣後約 US$3,080（約 NT$98,560 + 運費）',
    '建議零售定價 NT$360 / 支；給頂尖等批發客九折（NT$324）',
    '毛利約 10%，但可預收款項再出貨，無庫存風險',
    '待辦：與頂尖確認是否接受九折 NT$324 / 支，確認後再決定是否下單',
]:
    body_para(item, bullet=True)

heading('■ 網站二手工具欄位', level=2)
for item in [
    '網站目前欄位已滿（WooCommerce / WarpLayers 限制）',
    '決議：將「設備」欄位併入「配件」，「收縮機 PTR」改標記為配件類',
    '待辦：通知大砲（網站管理員）執行合併，空出欄位給「二手工具」',
]:
    body_para(item, bullet=True)

heading('■ KeyLand 膠水補貨', level=2)
body_para('已下訂 20 包 × 1kg，等待 KeyLand 回覆中，若無回應再催促', bullet=True)

heading('■ 工作效率（內部提醒）', level=2)
for item in [
    '早餐訂購、分餐等非職責工作應委由店長安排，避免佔用正式上班時間',
    'Jay 要求Mary提升工作效率，清楚知道自己每日任務',
]:
    body_para(item, bullet=True)

# ══════════════════════════════════════════════════════════════════════════════
# 7. Action Items
# ══════════════════════════════════════════════════════════════════════════════
heading('7．行動清單 Action Items')
make_table(
    ['待辦事項', '負責人', '期限 / 狀態'],
    [
        ['提供 KECO 網站帳號密碼給採購同仁',          'Jay',          '待辦'],
        ['VIP / Brickwick / Anson Lisa 窗口資料補充', 'Jay',          '待辦'],
        ['與頂尖確認錘子九折報價（NT$324/支）',       'Jay',          '待辦'],
        ['催促 KeyLand 膠水訂單回覆',                  'Jay',          '待辦'],
        ['通知大砲合併「設備」至「配件」欄位',         'Jay',          '待辦'],
        ['Kenji 光板退款金額確認（20–30 RMB/片）',    'Jay',          '待辦'],
        ['建立「代結帳區」雲端共用表格',               'Jay ＋ 主任', '7/1 前完成'],
        ['主任報到日確認（7/1 入職）',                 'Jay',          '待確認'],
        ['主任入職後執行全倉盤點',                     '主任',         '7/1 入職後第一週'],
        ['微信公司手機交接，更新微信顯示名稱',         'Jay ＋ 主任', '7/1 入職時'],
        ['大陸倉售價與移倉清單交接給主任',             'Jay',          '7/1 入職時'],
        ['Altra 已被取代庫存下架並執行下價',           'Jay',          '待辦'],
    ],
    col_widths=[7.5, 3.2, 6.0]
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(fp, 'Dentstage 凹痕工廠 PDR Tools  ｜  www.dentstagetools.com.tw  ｜  紀錄整理：Myke AI  ｜  2026.06.22',
    color=RGBColor(0x88,0x88,0x88), size=9)

doc.save(OUT_PATH)
print(f'Done: {OUT_PATH}')
